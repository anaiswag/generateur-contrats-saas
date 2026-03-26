import streamlit as st
from docx import Document
import pandas as pd
import io
import os
import tempfile
from openai import OpenAI
from pdf2docx import Converter

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="ContractAI Premium", page_icon="⚖️", layout="wide")
st.title("⚖️ Générateur de Contrats Intelligent")
st.markdown("Sans balises complexes • Avec Tableaux Dynamiques • Règles Logiques IA")

API_KEY = os.environ.get("OPENAI_API_KEY")
if not API_KEY:
    st.error("🚨 Erreur : La clé API OpenAI n'est pas configurée sur Railway.")
    st.stop()

# --- 2. FONCTION IA (Le Cerveau Logique) ---
def deduire_responsable_avec_ia(montant, regle):
    """L'IA lit la règle et le montant, et retourne juste le nom du responsable"""
    client = OpenAI(api_key=API_KEY)
    prompt = f"""
    Voici la règle de l'entreprise : "{regle}"
    Le montant total du contrat actuel est de {montant} €.
    En te basant strictement sur la règle, qui doit être le responsable/signataire de ce contrat ?
    Réponds UNIQUEMENT par le prénom et nom de la personne, sans aucune phrase autour.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

# --- 3. FONCTION DE MODIFICATION WORD (La Magie) ---
def modifier_document(chemin_fichier, remplacements, df_tableau):
    doc = Document(chemin_fichier)
    
    # A. Remplacer le texte partout (paragraphes normaux)
    for para in doc.paragraphs:
        for ancien, nouveau in remplacements.items():
            if ancien in para.text:
                para.text = para.text.replace(ancien, str(nouveau))
                
    # B. Parcourir les tableaux pour remplacer le texte ET gérer le tableau dynamique
    for table in doc.tables:
        if len(table.rows) > 0:
            # On lit la première ligne pour voir si c'est notre tableau de prestations
            en_tetes = [cell.text.strip().lower() for cell in table.rows[0].cells]
            
            if "description" in en_tetes and "channel" in en_tetes and "date" in en_tetes:
                # BINGO ! C'est le tableau dynamique. 
                # On sauvegarde la ligne d'exemple (pour la supprimer après)
                ligne_exemple = table.rows[1] if len(table.rows) > 1 else None
                
                # On ajoute les nouvelles lignes venant du site web
                for index, ligne_donnees in df_tableau.iterrows():
                    nouvelle_ligne = table.add_row()
                    nouvelle_ligne.cells[0].text = str(ligne_donnees.get("Description", ""))
                    nouvelle_ligne.cells[1].text = str(ligne_donnees.get("Channel", ""))
                    nouvelle_ligne.cells[2].text = str(ligne_donnees.get("Date", ""))
                
                # On supprime la ligne d'exemple d'origine du modèle
                if ligne_exemple:
                    table._tbl.remove(ligne_exemple._tr)
            else:
                # Si c'est un tableau normal, on fait juste du Chercher/Remplacer de texte
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for ancien, nouveau in remplacements.items():
                                if ancien in para.text:
                                    para.text = para.text.replace(ancien, str(nouveau))
    
    # Sauvegarder en mémoire
    fichier_final = io.BytesIO()
    doc.save(fichier_final)
    fichier_final.seek(0)
    return fichier_final

# --- 4. INTERFACE UTILISATEUR ---
st.markdown("### 📂 Étape 1 : Importez votre modèle (PDF ou Word)")
fichier_upload = st.file_uploader("Le document ne nécessite aucune balise de code.", type=["docx", "pdf"])

if fichier_upload is not None:
    # --- CONVERSION PDF SI BESOIN ---
    chemin_travail = ""
    with st.spinner("Préparation du document..."):
        if fichier_upload.name.endswith('.pdf'):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                tmp_pdf.write(fichier_upload.read())
                pdf_path = tmp_pdf.name
            docx_path = pdf_path.replace(".pdf", ".docx")
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            chemin_travail = docx_path
        else:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                tmp_docx.write(fichier_upload.read())
                chemin_travail = tmp_docx.name

    st.success("Document prêt ! Remplissez les données ci-dessous.")

    col1, col2 = st.columns(2)
    
    # --- COLONNE 1 : LOGIQUE ET VARIABLES ---
    with col1:
        st.markdown("### 🧠 Logique et Variables")
        st.info("Dites au programme quel texte remplacer dans votre modèle.")
        ancien_client = st.text_input("Texte du client à remplacer (ex: 'NOM_CLIENT') :", "NOM_CLIENT")
        nouveau_client = st.text_input("Nouveau nom du client :", "")
        
        ancien_responsable = st.text_input("Texte du responsable à remplacer (ex: 'SIGNATAIRE') :", "SIGNATAIRE")
        
        st.markdown("#### ⚙️ Règle automatique (Montant)")
        montant_total = st.number_input("Montant total du contrat (€) :", value=0)
        regle_ia = st.text_area("Règle d'attribution :", value="Si le montant dépasse 10000€, la responsable est Mme Dubois. Sinon, le responsable est M. Martin.")

    # --- COLONNE 2 : LE TABLEAU DYNAMIQUE ---
    with col2:
        st.markdown("### 📋 Tableau des Prestations")
        st.write("Ajoutez ou supprimez des lignes. Elles seront injectées dans le Word.")
        
        # Création du tableau interactif
        donnees_depart = pd.DataFrame([{"Description": "", "Channel": "", "Date": ""}])
        
        tableau_interactif = st.data_editor(
            donnees_depart,
            num_rows="dynamic", # Permet d'ajouter/supprimer des lignes (Le bouton +)
            use_container_width=True
        )

    # --- BOUTON DE GÉNÉRATION ---
    st.markdown("---")
    if st.button("✨ Analyser et Générer le Contrat Final", use_container_width=True):
        
        with st.spinner("🤖 L'IA réfléchit au responsable..."):
            responsable_trouve = deduire_responsable_avec_ia(montant_total, regle_ia)
            st.success(f"L'IA a déduit que le responsable est : **{responsable_trouve}**")
            
        with st.spinner("📝 Création du contrat et des tableaux..."):
            # Dictionnaire des mots à chercher / remplacer
            mots_a_remplacer = {
                ancien_client: nouveau_client,
                ancien_responsable: responsable_trouve,
                "MONTANT_TOTAL": f"{montant_total} €" # On peut ajouter autant de variables que l'on veut
            }
            
            # Lancement de la modification
            document_final = modifier_document(chemin_travail, mots_a_remplacer, tableau_interactif)
            
            st.balloons()
            st.download_button(
                label="📥 TÉLÉCHARGER LE CONTRAT FINAL (.docx)",
                data=document_final,
                file_name=f"Contrat_{nouveau_client}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
